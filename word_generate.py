# Импортируем необходимую функцию из библиотеки
from docx import Document
import json
from pathlib import Path
from docx.shared import Pt
from datetime import datetime
from dateutil.relativedelta import relativedelta

file_path = "files\\Отчет_реализация_Шаблон.docx"
# Открываем и загружаем JSON-файл

def read_paragraphs_from_docx(file_path):
  """
  Открывает .docx файл и считывает из него каждый абзац.

  Args:
    file_path (str): Путь к вашему .docx файлу.
  """
  try:
    # Открываем документ
    document = Document(file_path)

    # Проверяем, есть ли в документе абзацы
    if not document.paragraphs:
      print("В документе не найдено абзацев.")
      return

    print("Абзацы в документе:")
    # Перебираем все абзацы в документе
    for i, paragraph in enumerate(document.paragraphs):
      # paragraph.text содержит текст абзаца
      # Пропускаем пустые строки, если не хотите их выводить
      if paragraph.text.strip():
        print(f"Абзац {i+1}: {paragraph.text}")

  except Exception as e:
    print(f"Произошла ошибка: {e}")

def find_person_folder():
    """Находит единственную папку внутри директории 'People'."""
    people_dir = Path("People")
    if not people_dir.is_dir():
        print(f"❌ Ошибка: Директория '{people_dir}' не найдена.")
        return None
    subdirectories = [d for d in people_dir.iterdir() if d.is_dir()]
    if len(subdirectories) == 1:
        return subdirectories[0]
    print(f"❌ Ошибка: Внутри '{people_dir}' должна быть ровно одна папка.")
    return None

with open(f'{find_person_folder()}\\Информация.json', 'r', encoding='utf-8') as file:
    info = json.load(file)

with open(f'{find_person_folder()}\\Канцелярские расходы.json', 'r', encoding='utf-8') as file:
    canc = json.load(file)


try:
    # Открываем документ
    document = Document(file_path)


    target_paragraph = document.paragraphs[0]
    run = target_paragraph.add_run(info["Арбитражный суд"])

    target_paragraph = document.paragraphs[1]
    run = target_paragraph.add_run(info["№ дела"])

    table = document.tables[0]
    cell = table.cell(0,0)
    cell.paragraphs[0].add_run(datetime.now().strftime("%d.%m.%Y")+" г.")

    table = document.tables[1]
    cell = table.cell(2, 0)
    cell.paragraphs[0].add_run(info["ФИО должника"])

    table = document.tables[2]
    cell = table.cell(0, 1)
    cell.paragraphs[0].add_run(info["Дата рождения"] + "\n" + info["Место рождения"])

    table = document.tables[2]
    cell = table.cell(1, 1)
    try:
        cell.paragraphs[0].add_run(info["Ранее имевшиеся ФИО"])
    except:
        cell.paragraphs[0].add_run("-")

    table = document.tables[2]
    cell = table.cell(2, 1)
    cell.paragraphs[0].add_run(info["Место жительства"])

    table = document.tables[2]
    cell = table.cell(3, 1)
    cell.paragraphs[0].add_run(info["ИНН"])

    table = document.tables[2]
    cell = table.cell(4, 1)
    cell.paragraphs[0].add_run(info["СНИЛС"])

    table = document.tables[2]
    cell = table.cell(6, 1)
    cell.paragraphs[0].add_run(info["Арбитражный суд"])

    table = document.tables[2]
    cell = table.cell(7, 1)
    cell.paragraphs[0].add_run(info["№ дела"])

    table = document.tables[2]
    cell = table.cell(8, 1)
    cell.paragraphs[0].add_run(info["Дата решения"])

    table = document.tables[2]
    cell = table.cell(9, 1)
    cell.paragraphs[0].add_run(info["Дата решения"])

    table = document.tables[7]
    cell = table.cell(2, 3)
    cell.paragraphs[0].add_run(info["Дата решения"] + "-" + datetime.now().strftime("%d.%m.%Y"))

    table = document.tables[16]
    cell = table.cell(0, 1)
    cell.paragraphs[0].add_run("№ " + info["Номер объявления Коммерсант"] + " стр. " + info["Страница"] + " № " + info["Газета"] + " от " + info["Дата"])

    table = document.tables[16]
    cell = table.cell(1, 1)
    cell.paragraphs[0].add_run("№ " + info["№ сообщения"] + " от " + info["Дата публикации"])

    table = document.tables[16]
    cell = table.cell(3, 1)
    cell.paragraphs[0].add_run((datetime.strptime(info["Дата"], "%d.%m.%Y") + relativedelta(months=2)).strftime("%d.%m.%Y"))

    table = document.tables[19]
    cell = table.cell(3, 4)
    cell.paragraphs[0].add_run(info["Сумма Коммерсант"].replace(".",","))

    table = document.tables[19]
    cell = table.cell(3, 5)
    cell.paragraphs[0].add_run(datetime.now().strftime("%d.%m.%Y"))

    table = document.tables[19]
    cell = table.cell(3, 6)
    cell.paragraphs[0].add_run(info["Сумма Коммерсант"].replace(".",","))

    table = document.tables[19]
    cell = table.cell(4, 4)
    cell.paragraphs[0].add_run(info["Сумма ЕФРСБ"].replace(".",","))

    table = document.tables[19]
    cell = table.cell(4, 5)
    cell.paragraphs[0].add_run(datetime.now().strftime("%d.%m.%Y"))

    table = document.tables[19]
    cell = table.cell(4, 6)
    cell.paragraphs[0].add_run(info["Сумма ЕФРСБ"].replace(".",","))



    table = document.tables[19]
    cell = table.cell(5, 5)
    cell.paragraphs[0].add_run(datetime.now().strftime("%d.%m.%Y"))




    table = document.tables[19]
    cell = table.cell(6, 4)
    cell.paragraphs[0].add_run(canc["sum chancellery"].replace(".", ","))

    table = document.tables[19]
    cell = table.cell(6, 5)
    cell.paragraphs[0].add_run(datetime.now().strftime("%d.%m.%Y"))

    table = document.tables[19]
    cell = table.cell(6, 6)
    cell.paragraphs[0].add_run(canc["sum chancellery"].replace(".", ","))

    table = document.tables[19]
    cell = table.cell(7, 5)
    cell.paragraphs[0].add_run(datetime.now().strftime("%d.%m.%Y"))

    table = document.tables[19]
    cell = table.cell(6, 5)
    cell.paragraphs[0].add_run(datetime.now().strftime("%d.%m.%Y"))

    table = document.tables[19]
    cell = table.cell(9, 3)
    cell.paragraphs[0].add_run(str(float(info["Сумма ЕФРСБ"]) + float(info["Сумма Коммерсант"]) + float(canc["sum chancellery"]) + 26061.76).replace(".",","))

    table = document.tables[19]
    cell = table.cell(9, 6)
    cell.paragraphs[0].add_run(str(float(info["Сумма ЕФРСБ"]) + float(info["Сумма Коммерсант"]) + float(canc["sum chancellery"]) + 1061.76).replace(".", ","))

    # # Можно также задать форматирование для добавленного текста
    # run.bold = True  # Сделать текст жирным
    # run.italic = True # Сделать текст курсивным
    # font = run.font
    # font.name = 'Calibri'
    # font.size = Pt(12)

    # Сохраняем изменения в новый файл, чтобы не затереть оригинал
    document.save(f"{find_person_folder()}\\Отчет реализация {find_person_folder().name}.docx")
    print("Текст успешно добавлен и сохранен в новый файл.")


except Exception as e:
  print(f"Произошла ошибка: {e}")